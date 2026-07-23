from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name('glocalizer-backend-redesign-report.docx')

INK = '172033'
MUTED = '667085'
BLUE = '2E74B5'
DARK_BLUE = '1F4D78'
PALE_BLUE = 'E8EEF5'
PALE_GRAY = 'F4F6F9'
RISK_RED = '9B1C1C'
GREEN = '157347'


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill)
    tc_pr.append(shading)


def set_cell_margins(cell) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in('w:tcMar')
    if margins is None:
        margins = OxmlElement('w:tcMar')
        tc_pr.append(margins)
    for side in ('top', 'start', 'bottom', 'end'):
        node = margins.find(qn(f'w:{side}'))
        if node is None:
            node = OxmlElement(f'w:{side}')
            margins.append(node)
        node.set(qn('w:w'), '120' if side in ('start', 'end') else '90')
        node.set(qn('w:type'), 'dxa')


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_pr = table._tbl.tblPr
    width = OxmlElement('w:tblW')
    width.set(qn('w:w'), str(sum(widths)))
    width.set(qn('w:type'), 'dxa')
    table_pr.append(width)
    indent = OxmlElement('w:tblInd')
    indent.set(qn('w:w'), '120')
    indent.set(qn('w:type'), 'dxa')
    table_pr.append(indent)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(widths[index]))
            tc_w.set(qn('w:type'), 'dxa')
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run(run, size=None, color=None, bold=None) -> None:
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Malgun Gothic')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Malgun Gothic')
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def add_para(doc, text='', style='Normal', before=None, after=None, align=None, color=None, size=None, bold=False):
    paragraph = doc.add_paragraph(style=style)
    if before is not None:
        paragraph.paragraph_format.space_before = Pt(before)
    if after is not None:
        paragraph.paragraph_format.space_after = Pt(after)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run(run, size=size, color=color, bold=bold)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.2
        set_run(p.add_run(item), size=10.5, color=INK)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f'Heading {level}')
    set_run(paragraph.add_run(text), color=BLUE if level < 3 else DARK_BLUE, bold=True)
    return paragraph


def add_callout(doc, label, text, color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_BLUE if color == BLUE else PALE_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f'{label}  ')
    set_run(run, size=10.5, color=color, bold=True)
    run = p.add_run(text)
    set_run(run, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_geometry(table, widths)
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, PALE_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(value), size=9.5, color=DARK_BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(value), size=9.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles['Normal']
    normal.font.name = 'Malgun Gothic'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for level, size, before, after, color in [(1, 16, 16, 8, BLUE), (2, 13, 12, 6, BLUE), (3, 12, 8, 4, DARK_BLUE)]:
        style = doc.styles[f'Heading {level}']
        style.font.name = 'Malgun Gothic'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.add_run('Glocalizer | Backend Redesign Report'), size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(footer.add_run('Internal technical report · 2026-07-23'), size=8.5, color=MUTED)


def build_report():
    doc = Document()
    configure_document(doc)

    add_para(doc, 'TECHNICAL DECISION REPORT', before=18, after=6, align=WD_ALIGN_PARAGRAPH.CENTER, size=10, color=BLUE, bold=True)
    add_para(doc, 'Glocalizer 백엔드 재구성 보고서', after=6, align=WD_ALIGN_PARAGRAPH.CENTER, size=24, color=INK, bold=True)
    add_para(doc, '무료 우선 OCR·번역 파이프라인과 안정적인 PNG 편집 흐름', after=22, align=WD_ALIGN_PARAGRAPH.CENTER, size=13, color=MUTED)

    add_matrix(doc, ['구분', '내용'], [
        ('대상', 'Glocalizer 이미지 현지화 backend'),
        ('작성일', str(date.today())),
        ('의사결정', 'Gemini Vision 단일 의존 구조를 분리형 OCR + translation provider 구조로 전환'),
        ('우선 범위', '이미지 1장 MVP를 안정화한 뒤 다중 asset과 ZIP으로 확장'),
    ], [2100, 7260])

    add_callout(doc, '권고안', 'OCR 좌표는 local PaddleOCR bridge가 담당하고, 번역은 text-only provider를 교체 가능하게 분리한다. 자동 cleanup과 Canvas export는 기존 역할을 유지한다.')

    add_heading(doc, '1. 현황과 문제 정의')
    add_para(doc, '현재 backend는 Gemini Vision 한 번의 요청에 이미지 OCR, 대표 영역 선택, 번역 후보, 좌표 JSON 생성을 모두 맡긴다. 이 구조는 간결하지만 provider의 모델 폐기, quota, 응답 형식 변화가 전체 작업 실패로 바로 이어진다.')
    add_matrix(doc, ['관찰된 문제', '영향', '구조적 원인'], [
        ('모델 폐기', '프로젝트가 즉시 failed 상태로 전환', 'GEMINI_MODEL이 pipeline과 failure code에 강하게 결합'),
        ('오류 정보 손실', 'Editor에는 동일한 일반 실패 문구만 표시', 'SDK의 HTTP status·provider message를 AppError가 덮어씀'),
        ('좌표 품질 변동', 'cleanup·editor·PNG 합성 위치가 흔들림', 'vision LLM이 OCR과 polygon을 동시에 추론'),
        ('확장 어려움', '다중 asset과 재시도 정책이 복잡해짐', '한 service가 download부터 DB 저장까지 수행'),
    ], [2200, 3000, 4160])

    add_heading(doc, '2. 목표 아키텍처', 1)
    add_para(doc, '새 구조는 역할을 세 계층으로 고정한다. API는 job 생성과 결과 조회만 담당하고, worker는 asset 상태 전이만 orchestration하며, provider는 OCR 또는 번역 하나만 책임진다.')
    add_matrix(doc, ['단계', '책임', '입력 / 출력'], [
        ('API', 'signed upload, project token 인증, job enqueue, 상태 조회', 'project·asset request / job id·status'),
        ('Worker', 'asset별 순서·재시도·실패 기록·project aggregate', 'queued asset / completed 또는 failed asset'),
        ('OCR engine', '한글 text, polygon, confidence 추출', 'preprocessed image / normalized region[]'),
        ('Translation provider', '원문과 context 기반 후보·추천 스타일 생성', 'source text / candidates[]'),
        ('Cleanup', 'transparent 또는 solid cleanup 결과 생성', 'image + primary region / cleaned image 또는 manual fallback'),
    ], [1500, 3900, 3960])

    add_callout(doc, '프로세스 경계', 'PaddleOCR는 별도 HTTP 서버가 아니라 backend worker가 관리하는 persistent Python child process로 실행한다. 개발 환경에서도 frontend와 backend 두 서버만 열면 된다.', GREEN)

    add_heading(doc, '3. 처리 흐름과 상태 전이')
    add_para(doc, '상태는 project가 아니라 asset을 진실의 원천으로 삼는다. project progress는 asset 상태를 집계해 만든다. 각 단계의 provider·model·latency·원본 오류는 processing run으로 남긴다.')
    add_matrix(doc, ['순서', 'asset 상태', '실패 시 처리'], [
        ('1', 'uploaded → preprocessing', '파일 형식·크기·이미지 decode 오류를 terminal failed로 기록'),
        ('2', 'recognizing', 'OCR bridge timeout은 retryable failed; text 없음은 OCR_TEXT_NOT_FOUND'),
        ('3', 'translating', 'provider 401/429/5xx/JSON 오류를 분리해 기록하고 retry policy 적용'),
        ('4', 'cleaning', '복잡한 배경은 failed가 아니라 needsManualCleanup=true로 completed'),
        ('5', 'completed', 'editor가 original 또는 cleaned base 위에 동일한 normalized state를 복원'),
    ], [800, 2800, 5760])

    add_heading(doc, '4. Provider 전략')
    add_matrix(doc, ['역할', 'Primary', 'Fallback / 정책'], [
        ('OCR + bbox', 'PaddleOCR local bridge', '다른 cloud vision으로 자동 전송하지 않음'),
        ('번역 후보', 'Groq text provider', 'rate limit은 retryable; 자동 OpenRouter fallback 금지'),
        ('오프라인 번역', '선택형 local NLLB provider', 'GPU/속도 벤치마크 통과 시 활성화'),
        ('자동 cleanup', 'Sharp 기반 existing implementation', 'low quality는 manual cleanup으로 완료'),
    ], [1800, 3000, 4560])
    add_para(doc, 'OpenRouter 무료 router는 모델과 데이터 경로가 요청마다 달라질 수 있어 기본 자동 fallback으로 사용하지 않는다. 필요하면 관리자가 명시적으로 켜는 development-only provider로 둔다.', color=MUTED, size=9.5)

    add_heading(doc, '5. 코드·데이터 구조')
    add_matrix(doc, ['영역', '신규 책임', '핵심 파일 예시'], [
        ('pipeline', '단계 orchestration과 asset 상태 전이', 'pipeline/asset-pipeline.ts'),
        ('ocr', 'provider interface, Python bridge, normalized polygon', 'ocr/paddle/paddle-bridge.ts'),
        ('translation', 'Groq·local provider, 후보 검증, model metadata', 'translation/translation-provider.ts'),
        ('observability', 'provider/model/latency/error classification 기록', 'observability/processing-run.repository.ts'),
        ('worker', 'job claim, stale recovery, retryable failure 재큐잉', 'worker/job-runner.ts'),
    ], [1800, 3600, 3960])

    add_para(doc, 'DB는 기존 projects, assets, ocr_regions, translations, jobs, editor_states를 유지한다. numbered migration으로 processing_runs 테이블과 assets의 provider failure metadata만 추가한다. Gemini, GLM 등 provider 이름을 컬럼명·error code에 넣지 않는다.', after=10)

    add_heading(doc, '6. 구현 로드맵')
    add_matrix(doc, ['단계', '작업', '통과 기준'], [
        ('A. 분리', 'Gemini 전용 pipeline 제거, provider-neutral type·error code 도입', '기존 API contract와 unit test 유지'),
        ('B. OCR', 'persistent PaddleOCR bridge·polygon normalization 구현', '실제 이미지 bbox가 Canvas와 일치'),
        ('C. 번역', 'Groq text provider·JSON validator·provider health check 구현', '401/404/429/timeout이 사용자 메시지와 DB에 구분됨'),
        ('D. 관측', 'processing run·asset retry·stale job recovery 추가', '재시도 후에도 원인과 attempt 확인 가능'),
        ('E. 벤치마크', '실제 이모티콘 30장으로 OCR·번역·latency 측정', '품질·속도 기준 충족 후 multi-file 진행'),
    ], [900, 4300, 4160])

    add_heading(doc, '7. 수용 기준과 리스크')
    add_bullets(doc, [
        '이미지 1장 처리에서 한글 원문·bbox·번역 후보가 저장되고 editor 새로고침 후 동일하게 복원된다.',
        'PaddleOCR 실패, Groq rate limit, translation JSON 오류가 서로 다른 error code와 재시도 가능 여부로 기록된다.',
        '복잡한 배경은 작업 자체를 실패시키지 않고 manualCleanup 상태로 editor에 전달된다.',
        '다운로드 PNG는 editor Canvas와 같은 normalized rect, text 위치, 색상, 크기, 회전을 반영한다.',
        'provider API key는 backend env에만 존재하며 frontend bundle·로그·error response에는 노출되지 않는다.',
    ])
    add_callout(doc, '리스크 관리', 'PaddleOCR의 한국어 품질과 CPU/GPU latency는 먼저 30장 벤치마크로 측정한다. 기준 미달이면 OCR provider만 교체할 수 있도록 interface와 storage schema를 고정한다.', RISK_RED)

    add_heading(doc, '8. 참고 자료')
    add_para(doc, '• Google Gemini 모델 문서: https://ai.google.dev/gemini-api/docs/models', size=9.5, color=MUTED)
    add_para(doc, '• Groq Llama 4 Scout 문서: https://console.groq.com/docs/model/meta-llama/llama-4-scout-17b-16e-instruct', size=9.5, color=MUTED)
    add_para(doc, '• OpenRouter Free Router: https://openrouter.ai/openrouter/free', size=9.5, color=MUTED)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == '__main__':
    build_report()
